#!/usr/bin/env python3
"""
Word Document to HTML Converter for CIQMS
Converts .docx files to clean HTML suitable for web display
"""

import os
import sys
from pathlib import Path
import argparse


def install_requirements():
    """Install required packages if not available"""
    try:
        import mammoth
        import docx
    except ImportError:
        print("Installing required packages...")
        os.system("pip install python-docx mammoth")
        import mammoth
        import docx


def convert_with_mammoth(docx_path, output_path=None):
    """
    Convert Word document to HTML using mammoth (cleanest output)
    """
    try:
        import mammoth

        if output_path is None:
            output_path = docx_path.replace(".docx", ".html").replace(".doc", ".html")

        with open(docx_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html_content = result.value

            # Create a complete HTML document
            full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Converted Document</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 10px 0;
        }}
        table, th, td {{
            border: 1px solid #ddd;
        }}
        th, td {{
            padding: 8px;
            text-align: left;
        }}
        th {{
            background-color: #f2f2f2;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""

        with open(output_path, "w", encoding="utf-8") as html_file:
            html_file.write(full_html)

        print(f"✅ Successfully converted to: {output_path}")

        # Print any warnings
        if result.messages:
            print("\n⚠️  Conversion warnings:")
            for message in result.messages:
                print(f"   {message}")

        return output_path

    except Exception as e:
        print(f"❌ Error with mammoth conversion: {e}")
        return None


def convert_with_python_docx(docx_path, output_path=None):
    """
    Convert Word document to HTML using python-docx (basic conversion)
    """
    try:
        from docx import Document

        if output_path is None:
            output_path = docx_path.replace(".docx", ".html").replace(".doc", ".html")

        doc = Document(docx_path)

        html_content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Simple paragraph conversion
                html_content.append(f"<p>{paragraph.text}</p>")

        # Handle tables
        for table in doc.tables:
            html_content.append("<table>")
            for row in table.rows:
                html_content.append("<tr>")
                for cell in row.cells:
                    html_content.append(f"<td>{cell.text}</td>")
                html_content.append("</tr>")
            html_content.append("</table>")

        full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Converted Document</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 10px 0;
        }}
        table, th, td {{
            border: 1px solid #ddd;
        }}
        th, td {{
            padding: 8px;
            text-align: left;
        }}
    </style>
</head>
<body>
{''.join(html_content)}
</body>
</html>"""

        with open(output_path, "w", encoding="utf-8") as html_file:
            html_file.write(full_html)

        print(f"✅ Successfully converted to: {output_path}")
        return output_path

    except Exception as e:
        print(f"❌ Error with python-docx conversion: {e}")
        return None


def convert_document(input_path, output_path=None, method="mammoth"):
    """
    Main conversion function
    """
    if not os.path.exists(input_path):
        print(f"❌ File not found: {input_path}")
        return None

    if not input_path.lower().endswith((".docx", ".doc")):
        print(f"❌ File must be .doc or .docx format")
        return None

    print(f"📄 Converting: {input_path}")
    print(f"🔧 Method: {method}")

    if method == "mammoth":
        return convert_with_mammoth(input_path, output_path)
    elif method == "python-docx":
        return convert_with_python_docx(input_path, output_path)
    else:
        print(f"❌ Unknown method: {method}")
        return None


def main():
    parser = argparse.ArgumentParser(description="Convert Word documents to HTML")
    parser.add_argument("input", help="Path to Word document (.docx or .doc)")
    parser.add_argument("-o", "--output", help="Output HTML file path")
    parser.add_argument(
        "-m",
        "--method",
        choices=["mammoth", "python-docx"],
        default="mammoth",
        help="Conversion method (default: mammoth)",
    )
    parser.add_argument(
        "--install", action="store_true", help="Install required packages"
    )

    args = parser.parse_args()

    if args.install:
        install_requirements()
        return

    result = convert_document(args.input, args.output, args.method)

    if result:
        print(f"\n🎉 Conversion complete!")
        print(f"📂 HTML file: {result}")

        # Open in browser (optional)
        try:
            import webbrowser

            webbrowser.open(f"file://{os.path.abspath(result)}")
            print("🌐 Opened in browser")
        except:
            pass


if __name__ == "__main__":
    main()
